# Assemble final CUREUS_READY doc with figures, table, and reference renumbering.

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement

base = Path("/mnt/data")
in_path = base / "National_Trends_GPAD_fix_citation_order_sorted.docx"
doc = Document(str(in_path))

def find_heading_idx(d, heading):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().lower() == heading.lower():
            return i
    return None

def insert_paragraph_after(paragraph, text):
    new_p = paragraph._element.addnext(OxmlElement('w:p'))
    p_obj = paragraph._p.getnext()  # new w:p element
    # Create a docx paragraph object from the element
    from docx.text.paragraph import Paragraph
    from docx.oxml.text.paragraph import CT_P
    return Paragraph(p_obj, paragraph._parent)

# ---------- 1) Ensure Figures section contains images + captions ----------
figs_idx = find_heading_idx(doc, "Figures")
fig1 = base / "Figure_1_Total_Appointments.png"
fig2 = base / "Figure_2_Mode_Share.png"

if figs_idx is not None and fig1.exists() and fig2.exists():
    # Remove any existing figure captions/images immediately after heading (up to ~10 lines or until next heading)
    to_remove = []
    i = figs_idx + 1
    while i < len(doc.paragraphs) and i < figs_idx + 20:
        t = doc.paragraphs[i].text.strip()
        if t.lower().startswith("figure ") or t == "":
            to_remove.append(i)
            i += 1
            continue
        # Stop if we hit another major heading
        if t.isupper() and len(t) < 60:
            break
        # Also remove non-empty stray text under Figures
        to_remove.append(i)
        i += 1
    for idx in reversed(to_remove):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    # Insert Figure 1
    p_after = doc.paragraphs[figs_idx]
    p_img1 = p_after.insert_paragraph_after("")
    run1 = p_img1.add_run()
    run1.add_picture(str(fig1), width=Inches(6))
    p_cap1 = p_img1.insert_paragraph_after("Figure 1. Total GP Appointments in England (Oct 2021–May 2025). Source: NHS England, Appointments in General Practice (National Overview CSVs).")

    # Insert Figure 2
    p_img2 = p_cap1.insert_paragraph_after("")
    run2 = p_img2.add_run()
    run2.add_picture(str(fig2), width=Inches(6))
    p_cap2 = p_img2.insert_paragraph_after("Figure 2. Mode of GP Appointments in England (Oct 2021–May 2025). Shares exclude 'Unknown/Other' from the denominator; minor differences may appear versus NHS Key Facts. Source: NHS England, Appointments in General Practice (National Overview CSVs).")

# ---------- 2) Ensure Table 1 is a proper Word table ----------
tbl_head_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower().startswith("table 1 — key milestones"):
        tbl_head_idx = i
        # Ensure heading text is the 2021–2025 version
        doc.paragraphs[i].text = "Table 1 — Key Milestones and Observations (2021–2025)"
        break

if tbl_head_idx is not None:
    # Remove lines under the table heading until a blank line
    j = tbl_head_idx + 1
    to_remove = []
    while j < len(doc.paragraphs):
        t = doc.paragraphs[j].text.strip()
        to_remove.append(j)
        if t == "":
            break
        j += 1
    for idx in reversed(to_remove):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    # Insert the table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Event/Observation"
    hdr[2].text = "Mode Share"
    hdr[3].text = "Notes"
    rows = [
        ("Oct 2021", "Post-acute stabilisation; hybrid model established", "≈66.9% face-to-face", "CSV-derived, primary denominator"),
        ("Jun 2022", "Hybrid sustained; telephone dominant within remote", "≈67.1% face-to-face", "CSV-derived, primary denominator"),
        ("Dec 2023", "Post-pandemic plateau", "≈69.2% face-to-face", "CSV-derived, primary denominator"),
        ("Mar 2024", "Plateau peak (Key facts)", "65.4% face-to-face", "NHS Key Facts"),
        ("Apr 2025", "Latest window (Key facts)", "63.8% face-to-face", "NHS Key Facts"),
        ("May 2025", "Latest reported month (Key facts)", "63.5% face-to-face", "NHS Key Facts"),
    ]
    for r in rows:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = r

# ---------- 3) Renumber references by first-mention order & rebuild list ----------
def find_heading_idx(d, heading):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().lower() == heading.lower():
            return i
    return None

refs_idx = find_heading_idx(doc, "References")
body_limit = refs_idx if refs_idx is not None else len(doc.paragraphs)

pat = re.compile(r"\[([0-9,\s–\-]+)\]")
def expand_group(s):
    nums = []
    for part in [x.strip() for x in s.split(",")]:
        if re.search(r"[–\-]", part):
            ab = [x for x in re.split(r"[–\-]", part) if x.strip().isdigit()]
            if len(ab) == 2:
                a, b = map(int, ab)
                nums.extend(range(a, b + 1))
        elif part.isdigit():
            nums.append(int(part))
    return nums

# First-mention order
seen = set()
order = []
for p in doc.paragraphs[:body_limit]:
    for m in pat.finditer(p.text):
        nums = expand_group(m.group(1))
        for n in nums:
            if n not in seen:
                seen.add(n); order.append(n)

# Parse existing Reference entries
entries = []
curr = None; buf = ""
for p in doc.paragraphs[refs_idx+1:]:
    t = p.text.strip()
    if not t:
        continue
    m = re.match(r"^(\d+)\.\s*(.*)", t)
    if m:
        if curr is not None:
            entries.append((curr, buf.strip()))
        curr = int(m.group(1))
        buf = m.group(2)
    else:
        if curr is not None:
            buf += " " + t
if curr is not None:
    entries.append((curr, buf.strip()))

old_to_text = {n: txt for n, txt in entries}

# Map old->new
mapper = {old: i+1 for i, old in enumerate(order)}

# Replace in-body citations with mapped numbers and sorted/collapsed runs
def collapse_mapped(nums):
    mapped = sorted(set(mapper.get(n, n) for n in nums))
    # Build runs
    runs = []
    start = prev = None
    for n in mapped:
        if start is None:
            start = prev = n
        elif n == prev + 1:
            prev = n
        else:
            runs.append((start, prev))
            start = prev = n
    if start is not None:
        runs.append((start, prev))
    parts = []
    for a, b in runs:
        if a == b:
            parts.append(str(a))
        elif b == a + 1:
            parts.append(str(a)); parts.append(str(b))
        else:
            parts.append(f"{a}–{b}")
    return "[" + ",".join(parts) + "]"

for p in doc.paragraphs[:body_limit]:
    def repl(m):
        nums = expand_group(m.group(1))
        return collapse_mapped(nums)
    p.text = pat.sub(repl, p.text)

# Rebuild the references section in the new order
# Remove existing lines after the "References" heading
for i in range(len(doc.paragraphs)-1, refs_idx, -1):
    el = doc.paragraphs[i]._element
    el.getparent().remove(el)

# Insert new references
anchor = doc.paragraphs[refs_idx]._element
for new_num, old in enumerate(order, start=1):
    txt = old_to_text.get(old, "")
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = f"{new_num}. {txt}"
    r.append(t); p.append(r)
    anchor.addnext(p)

# ---------- 4) Save ----------
out_path = base / "CUREUS_READY_Telemedicine_vs_InPerson_NHS_2021-2025.docx"
doc.save(str(out_path))
str(out_path)
